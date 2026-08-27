"""
Gera um documento Word (.docx) com o roteiro de teste de Primeiro Emplacamento,
a partir dos dados estruturados em data/roteiros_teste.py.

Inclui: capa, ambiente, cada cenario com passo a passo (tabela com acao, camada,
transacao e resultado esperado) e a legenda de transacoes.

Saida: docs/Roteiro_de_Teste_Primeiro_Emplacamento.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from data.roteiros_teste import get_roteiros, get_transacoes, AMBIENTE

PRETO = RGBColor(0x1A, 0x1A, 0x1A)
CINZA = RGBColor(0x66, 0x66, 0x66)


def add_titulo(doc, texto, size=16, cor=PRETO, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = cor
    return p


def main():
    doc = Document()

    # estilo base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ---- Capa ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Roteiro de Teste')
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRETO
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run('Primeiro Emplacamento de Veiculo Zero-KM')
    rs.font.size = Pt(14)
    rs.font.color.rgb = CINZA
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Prodesp / DETRAN-SP  |  BRQ Digital Solutions').font.color.rgb = CINZA

    # ---- Ambiente ----
    add_titulo(doc, 'Ambiente de execucao', size=13)
    p = doc.add_paragraph()
    p.add_run('Plataforma: ').bold = True
    p.add_run(f"{AMBIENTE['plataforma']}\n")
    p.add_run('Host: ').bold = True
    p.add_run(f"{AMBIENTE['host']}   ")
    p.add_run('IP: ').bold = True
    p.add_run(f"{AMBIENTE['ip']}   ")
    p.add_run('Window: ').bold = True
    p.add_run(AMBIENTE['window'])

    # ---- Legenda de transacoes ----
    add_titulo(doc, 'Transacoes utilizadas', size=13)
    trans = get_transacoes()
    tab = doc.add_table(rows=1, cols=3)
    tab.style = 'Light Grid Accent 1'
    hdr = tab.rows[0].cells
    hdr[0].text = 'Transacao'; hdr[1].text = 'Descricao'; hdr[2].text = 'Tela / Programa'
    for cod, info in trans.items():
        row = tab.add_row().cells
        row[0].text = cod
        row[1].text = info.get('descricao', '')
        extra = info.get('tela', '')
        if info.get('programa'):
            extra = (extra + ' | ' if extra else '') + f"{info['programa']} ({info.get('fonte','')})"
        row[2].text = extra

    # ---- Cenarios ----
    roteiros = get_roteiros()
    imgs_base = Path('frontend/static/roteiros')
    for rot in roteiros:
        doc.add_page_break()
        add_titulo(doc, rot['nome'], size=15)
        dt = rot['dados_teste']
        ident = (f"CPF {dt['cpf']}" if 'cpf' in dt else f"CNPJ {dt['cnpj']}")
        p = doc.add_paragraph()
        p.add_run('Categoria: ').bold = True
        p.add_run(f"{rot['categoria']}    ")
        p.add_run('Chassi: ').bold = True
        p.add_run(f"{dt['chassi']}    ")
        p.add_run(ident.split()[0] + ': ').bold = True
        p.add_run(ident.split()[1])

        # tabela de passos
        tab = doc.add_table(rows=1, cols=5)
        tab.style = 'Light Grid Accent 1'
        h = tab.rows[0].cells
        h[0].text = '#'; h[1].text = 'Acao'; h[2].text = 'Camada'
        h[3].text = 'Transacao'; h[4].text = 'Resultado esperado'
        for passo in rot['passos']:
            c = tab.add_row().cells
            c[0].text = str(passo['ordem'])
            c[1].text = passo['titulo'] + (('\n' + passo['descricao']) if passo.get('descricao') else '')
            c[2].text = passo.get('camada', '')
            c[3].text = passo.get('transacao') or '-'
            c[4].text = passo.get('resultado_esperado', '')

        # anexar algumas telas de referencia do cenario
        pasta = imgs_base / rot['id']
        if pasta.exists():
            telas = sorted(pasta.glob('*.png'))
            # amostra: primeira, meio e ultima para nao inchar o doc
            amostra_idx = sorted(set([0, len(telas)//3, 2*len(telas)//3, len(telas)-1]))
            add_titulo(doc, 'Telas de referencia', size=12, space_before=10)
            for i in amostra_idx:
                if 0 <= i < len(telas):
                    try:
                        doc.add_picture(str(telas[i]), width=Inches(5.5))
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        rc = cap.add_run(f'Tela {i+1} de {len(telas)}')
                        rc.font.size = Pt(8)
                        rc.font.color.rgb = CINZA
                    except Exception:
                        pass

    out = Path('docs/Roteiro_de_Teste_Primeiro_Emplacamento.docx')
    out.parent.mkdir(exist_ok=True)
    doc.save(str(out))
    print('Documento gerado:', out)


if __name__ == '__main__':
    main()
